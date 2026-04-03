"""Comprehensive tests for document editing skills (Word + Excel + utility)."""

import io
import pytest

from document_skills import (
    create_word,
    read_word,
    edit_word,
    create_excel,
    read_excel,
    edit_excel,
    detect_format,
)


# ── Word (.docx) ────────────────────────────────────────────────────────


async def test_create_word_basic():
    """Create a doc with title + content, verify valid .docx bytes."""
    result = await create_word("Test Title", "Hello world\nSecond paragraph")
    assert isinstance(result, bytes)
    assert len(result) > 100
    assert result[:2] == b"PK"  # .docx is a ZIP archive


async def test_read_word_extracts_text():
    """Create → read round-trip preserves content."""
    original = "The quick brown fox jumps over the lazy dog"
    doc_bytes = await create_word("Roundtrip", original)
    extracted = await read_word(doc_bytes)
    assert original in extracted


async def test_read_word_with_tables():
    """Tables embedded in a doc are extracted as pipe-delimited text."""
    from docx import Document

    doc = Document()
    doc.add_heading("Table Doc", level=1)
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Name"
    table.cell(0, 1).text = "Score"
    table.cell(1, 0).text = "Alice"
    table.cell(1, 1).text = "95"

    buf = io.BytesIO()
    doc.save(buf)

    text = await read_word(buf.getvalue())
    assert "Name" in text
    assert "Score" in text
    assert "Alice" in text
    assert "95" in text


async def test_edit_word_find_replace():
    """Single find-and-replace in a paragraph."""
    doc_bytes = await create_word("Edit Test", "Hello Alice")
    edited = await edit_word(doc_bytes, {"Hello": "Goodbye"})
    text = await read_word(edited)
    assert "Goodbye" in text
    assert "Hello" not in text


async def test_edit_word_multiple_replacements():
    """Multiple find-replace pairs applied in one call."""
    doc_bytes = await create_word("Multi", "Hello Alice, welcome to Q1 review")
    edited = await edit_word(doc_bytes, {"Alice": "Bob", "Q1": "Q2"})
    text = await read_word(edited)
    assert "Bob" in text
    assert "Q2" in text
    assert "Alice" not in text
    assert "Q1" not in text


async def test_edit_word_no_match():
    """Edit with text that doesn't exist leaves file unchanged."""
    original_content = "Nothing to change here"
    doc_bytes = await create_word("NoMatch", original_content)
    edited = await edit_word(doc_bytes, {"NONEXISTENT": "replacement"})
    text = await read_word(edited)
    assert original_content in text


async def test_read_word_empty():
    """Empty doc produces minimal or empty text."""
    doc_bytes = await create_word("Empty", "")
    text = await read_word(doc_bytes)
    # Title heading is still present
    assert "Empty" in text
    # No body paragraphs beyond the title
    lines = [l for l in text.split("\n") if l.strip() and l.strip() != "Empty"]
    assert len(lines) == 0


# ── Excel (.xlsx) ────────────────────────────────────────────────────────


async def test_create_excel_basic():
    """Create with headers + rows, verify valid .xlsx bytes."""
    result = await create_excel(
        "Test Sheet",
        headers=["Name", "Age"],
        rows=[["Alice", "30"], ["Bob", "25"]],
    )
    assert isinstance(result, bytes)
    assert len(result) > 100
    assert result[:2] == b"PK"


async def test_read_excel_formats_table():
    """Create → read returns markdown-style table with headers and separator."""
    xlsx = await create_excel(
        "Sales",
        headers=["Product", "Revenue"],
        rows=[["Widget", "1000"], ["Gadget", "2000"]],
    )
    text = await read_excel(xlsx)
    assert "Product" in text
    assert "Revenue" in text
    assert "Widget" in text
    assert "---" in text  # markdown separator


async def test_read_excel_multiple_sheets():
    """Workbook with two sheets — both are read."""
    from openpyxl import Workbook

    wb = Workbook()
    ws1 = wb.active
    ws1.title = "Alpha"
    ws1.append(["X", "Y"])
    ws1.append([1, 2])

    ws2 = wb.create_sheet("Beta")
    ws2.append(["A", "B"])
    ws2.append([3, 4])

    buf = io.BytesIO()
    wb.save(buf)
    wb.close()

    text = await read_excel(buf.getvalue())
    assert "Alpha" in text
    assert "Beta" in text
    assert "X" in text
    assert "A" in text


async def test_edit_excel_set_cell():
    """Edit cell A1 and verify the change is read back."""
    xlsx = await create_excel("Data", headers=["Col1"], rows=[["original"]])
    edited = await edit_excel(xlsx, [{"cell": "A1", "value": "Updated"}])
    text = await read_excel(edited)
    assert "Updated" in text


async def test_edit_excel_numeric_conversion():
    """Numeric string '42' is stored as int, not string."""
    from openpyxl import load_workbook

    xlsx = await create_excel("Nums", headers=["Val"], rows=[["0"]])
    edited = await edit_excel(xlsx, [{"cell": "A2", "value": "42"}])

    wb = load_workbook(io.BytesIO(edited))
    ws = wb.active
    assert ws["A2"].value == 42
    assert isinstance(ws["A2"].value, int)
    wb.close()


async def test_edit_excel_multiple_cells():
    """Multiple cell edits in one call."""
    xlsx = await create_excel("Multi", headers=["A", "B"], rows=[["x", "y"]])
    edited = await edit_excel(
        xlsx,
        [
            {"cell": "A2", "value": "alpha"},
            {"cell": "B2", "value": "beta"},
        ],
    )
    text = await read_excel(edited)
    assert "alpha" in text
    assert "beta" in text


async def test_edit_excel_nonexistent_sheet():
    """Edit targeting a bad sheet name falls back to active sheet."""
    xlsx = await create_excel("Fallback", headers=["H1"], rows=[["v1"]])
    edited = await edit_excel(
        xlsx, [{"sheet": "DoesNotExist", "cell": "A2", "value": "fallback_val"}]
    )
    text = await read_excel(edited)
    assert "fallback_val" in text


# ── Utility ──────────────────────────────────────────────────────────────


@pytest.mark.parametrize(
    "filename, expected",
    [
        ("report.docx", "word"),
        ("data.xlsx", "excel"),
        ("doc.pdf", "pdf"),
        ("image.png", None),
    ],
    ids=["word", "excel", "pdf", "unknown"],
)
def test_detect_format(filename, expected):
    assert detect_format(filename) == expected
