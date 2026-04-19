"""Document editing skills — read, edit, and create Word (.docx) and Excel (.xlsx) files."""
import io
import logging
from pathlib import Path

log = logging.getLogger(__name__)


# ── Word (.docx) ─────────────────────────────────────────────────────────

async def read_word(file_bytes: bytes) -> str:
    """Extract all text from a .docx file. Returns plain text."""
    import asyncio

    from docx import Document

    def _extract():
        doc = Document(io.BytesIO(file_bytes))
        parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text)
        # Also extract tables
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                parts.append(" | ".join(cells))
        return "\n".join(parts)

    return await asyncio.to_thread(_extract)


async def edit_word(file_bytes: bytes, edits: dict[str, str]) -> bytes:
    """Apply find-and-replace edits to a .docx file.

    Args:
        file_bytes: Original .docx content
        edits: Dict of {find_text: replace_text} pairs

    Returns:
        Modified .docx as bytes
    """
    import asyncio

    from docx import Document

    def _apply_edits():
        doc = Document(io.BytesIO(file_bytes))
        for para in doc.paragraphs:
            for find_text, replace_text in edits.items():
                if find_text in para.text:
                    for run in para.runs:
                        if find_text in run.text:
                            run.text = run.text.replace(find_text, replace_text)
        # Also edit in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for find_text, replace_text in edits.items():
                            if find_text in para.text:
                                for run in para.runs:
                                    if find_text in run.text:
                                        run.text = run.text.replace(find_text, replace_text)
        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()

    return await asyncio.to_thread(_apply_edits)


async def create_word(title: str, content: str, headers: list[str] | None = None) -> bytes:
    """Create a new .docx file with the given content.

    Args:
        title: Document title (added as Heading 1)
        content: Body text (paragraphs separated by newlines)
        headers: Optional list of section headers (added as Heading 2)

    Returns:
        .docx file as bytes
    """
    import asyncio

    from docx import Document

    def _create():
        doc = Document()
        doc.add_heading(title, level=1)

        paragraphs = content.split("\n")
        header_idx = 0

        for para_text in paragraphs:
            para_text = para_text.strip()
            if not para_text:
                continue

            if headers and header_idx < len(headers) and para_text == headers[header_idx]:
                doc.add_heading(para_text, level=2)
                header_idx += 1
            else:
                doc.add_paragraph(para_text)

        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()

    return await asyncio.to_thread(_create)


# ── Excel (.xlsx) ────────────────────────────────────────────────────────

async def read_excel(file_bytes: bytes, sheet_name: str | None = None) -> str:
    """Extract data from an .xlsx file as formatted text.

    Returns a text representation with headers and rows.
    """
    import asyncio

    from openpyxl import load_workbook

    def _extract():
        wb = load_workbook(io.BytesIO(file_bytes), read_only=True, data_only=True)
        parts = []

        sheets = [wb[sheet_name]] if sheet_name and sheet_name in wb.sheetnames else wb.worksheets

        for ws in sheets:
            parts.append(f"## Sheet: {ws.title}")
            rows_data = []
            for row in ws.iter_rows(values_only=True):
                row_strs = [str(cell) if cell is not None else "" for cell in row]
                rows_data.append(row_strs)

            if rows_data:
                header = rows_data[0]
                parts.append(" | ".join(header))
                parts.append(" | ".join(["---"] * len(header)))
                for row in rows_data[1:]:
                    while len(row) < len(header):
                        row.append("")
                    parts.append(" | ".join(row[:len(header)]))
            parts.append("")

        wb.close()
        return "\n".join(parts)

    return await asyncio.to_thread(_extract)


async def edit_excel(file_bytes: bytes, edits: list[dict]) -> bytes:
    """Apply cell edits to an .xlsx file.

    Args:
        file_bytes: Original .xlsx content
        edits: List of dicts with keys: sheet (optional), cell, value
               e.g. [{"cell": "A1", "value": "New Title"}, {"sheet": "Sheet2", "cell": "B3", "value": "42"}]

    Returns:
        Modified .xlsx as bytes
    """
    import asyncio

    from openpyxl import load_workbook

    def _apply_edits():
        wb = load_workbook(io.BytesIO(file_bytes))

        for edit in edits:
            sheet_name = edit.get("sheet", wb.active.title)
            cell_ref = edit["cell"]
            value = edit["value"]

            if sheet_name in wb.sheetnames:
                ws = wb[sheet_name]
            else:
                ws = wb.active

            try:
                value = float(value)
                if value == int(value):
                    value = int(value)
            except (ValueError, TypeError):
                pass

            ws[cell_ref] = value

        buf = io.BytesIO()
        wb.save(buf)
        wb.close()
        return buf.getvalue()

    return await asyncio.to_thread(_apply_edits)


async def create_excel(title: str, headers: list[str], rows: list[list], sheet_name: str = "Sheet1") -> bytes:
    """Create a new .xlsx file with headers and data rows.

    Args:
        title: Used as filename context (not in spreadsheet itself)
        headers: Column header names
        rows: List of row data (list of lists)
        sheet_name: Name for the worksheet

    Returns:
        .xlsx file as bytes
    """
    import asyncio

    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill

    def _create():
        wb = Workbook()
        ws = wb.active
        ws.title = sheet_name

        header_fill = PatternFill(start_color="4472C4", end_color="4472C4", fill_type="solid")
        header_font_white = Font(bold=True, color="FFFFFF")

        for col_idx, header in enumerate(headers, 1):
            cell = ws.cell(row=1, column=col_idx, value=header)
            cell.font = header_font_white
            cell.fill = header_fill

        for row_idx, row_data in enumerate(rows, 2):
            for col_idx, value in enumerate(row_data, 1):
                try:
                    value = float(value)
                    if value == int(value):
                        value = int(value)
                except (ValueError, TypeError):
                    pass
                ws.cell(row=row_idx, column=col_idx, value=value)

        for col in ws.columns:
            max_length = max(len(str(cell.value or "")) for cell in col)
            ws.column_dimensions[col[0].column_letter].width = min(max_length + 2, 50)

        buf = io.BytesIO()
        wb.save(buf)
        wb.close()
        return buf.getvalue()

    return await asyncio.to_thread(_create)


# ── Utility ──────────────────────────────────────────────────────────────

def detect_format(filename: str) -> str | None:
    """Detect document format from filename extension."""
    ext = Path(filename).suffix.lower()
    if ext in (".docx", ".doc"):
        return "word"
    elif ext in (".xlsx", ".xls"):
        return "excel"
    elif ext == ".pdf":
        return "pdf"
    return None


SUPPORTED_FORMATS = {
    "word": [".docx"],
    "excel": [".xlsx"],
    "pdf": [".pdf"],
}
