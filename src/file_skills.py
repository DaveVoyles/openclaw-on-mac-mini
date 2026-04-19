"""Local file skills — read and list files from the shared /ai-files mount.

Supported formats:
  Text/code: .txt .md .rst .csv .json .jsonl .yaml .yml .toml .ini .cfg .conf
             .log .py .js .ts .sh .bash .html .xml .css .sql .r
  Rich docs:  .pdf  (via pypdf)
              .docx (via python-docx)
              .xlsx .xls (via openpyxl)
"""

import logging
from pathlib import Path

log = logging.getLogger(__name__)

AI_FILES_DIR = Path("/ai-files")
MAX_FILE_SIZE = 10 * 1024 * 1024  # 10 MB cap (rich docs can be larger)
MAX_TEXT_CHARS = 200_000  # cap extracted text sent to LLM

TEXT_EXTENSIONS = {
    ".txt",
    ".md",
    ".markdown",
    ".rst",
    ".csv",
    ".json",
    ".jsonl",
    ".yaml",
    ".yml",
    ".toml",
    ".ini",
    ".cfg",
    ".conf",
    ".log",
    ".py",
    ".js",
    ".ts",
    ".sh",
    ".bash",
    ".html",
    ".xml",
    ".css",
    ".sql",
    ".r",
    ".R",
}
PDF_EXTENSIONS = {".pdf"}
WORD_EXTENSIONS = {".docx"}
EXCEL_EXTENSIONS = {".xlsx", ".xls"}


def _is_safe_path(path: Path) -> bool:
    """Ensure path is inside AI_FILES_DIR (prevent traversal)."""
    try:
        path.resolve().relative_to(AI_FILES_DIR.resolve())
        return True
    except ValueError:
        return False


def _read_pdf(path: Path) -> str:
    """Extract text from a PDF using pypdf."""
    try:
        import pypdf

        reader = pypdf.PdfReader(str(path))
        pages = []
        for i, page in enumerate(reader.pages):
            text = page.extract_text() or ""
            if text.strip():
                pages.append(f"--- Page {i + 1} ---\n{text.strip()}")
        if not pages:
            return "(PDF contains no extractable text — may be scanned/image-only)"
        return "\n\n".join(pages)
    except Exception as e:
        return f"Error reading PDF: {e}"


def _read_docx(path: Path) -> str:
    """Extract text from a .docx file using python-docx."""
    try:
        from docx import Document

        doc = Document(str(path))
        parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells]
                parts.append(" | ".join(cells))
        return "\n".join(parts) if parts else "(Document is empty)"
    except Exception as e:
        return f"Error reading DOCX: {e}"


def _read_excel(path: Path) -> str:
    """Extract data from an .xlsx/.xls file using openpyxl."""
    try:
        import openpyxl

        wb = openpyxl.load_workbook(str(path), read_only=True, data_only=True)
        parts = []
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            rows = []
            for row in ws.iter_rows(values_only=True):
                if any(v is not None for v in row):
                    rows.append(" | ".join(str(v) if v is not None else "" for v in row))
            if rows:
                parts.append(f"=== Sheet: {sheet_name} ===\n" + "\n".join(rows))
        return "\n\n".join(parts) if parts else "(Spreadsheet is empty)"
    except Exception as e:
        return f"Error reading Excel: {e}"


async def list_local_files(directory: str = "/ai-files") -> str:
    """List files in the shared AI files directory (or a subdirectory of it).

    Args:
        directory: Path to list. Must be /ai-files or a subdirectory.

    Returns:
        Formatted list of files with sizes and types.
    """
    import asyncio

    target = Path(directory)
    if not _is_safe_path(target):
        return "Error: directory must be within /ai-files"

    def _list():
        if not target.exists():
            return f"Directory not found: {directory}"
        if not target.is_dir():
            return f"Not a directory: {directory}"

        entries = []
        for item in sorted(target.iterdir()):
            size = item.stat().st_size if item.is_file() else 0
            kind = "dir" if item.is_dir() else item.suffix.lstrip(".") or "file"
            size_str = f"{size:,} bytes" if item.is_file() else ""
            entries.append(f"{'/' if item.is_dir() else ' '} {item.name}  {kind}  {size_str}".rstrip())

        if not entries:
            return f"Directory is empty: {directory}"
        return f"Contents of {directory} ({len(entries)} items):\n" + "\n".join(entries)

    return await asyncio.to_thread(_list)


async def read_local_file(path: str) -> str:
    """Read the contents of a file from the shared /ai-files directory.

    Supports plain text, code, PDF, Word (.docx), and Excel (.xlsx/.xls).

    Args:
        path: File path, e.g. /ai-files/report.pdf or just report.pdf

    Returns:
        File contents as text, or an error message.
    """
    import asyncio

    target = Path(path)
    if not target.is_absolute():
        target = AI_FILES_DIR / path

    if not _is_safe_path(target):
        return "Error: path must be within /ai-files"

    def _read():
        if not target.exists():
            return f"File not found: {path}"
        if not target.is_file():
            return f"Not a file: {path}"

        size = target.stat().st_size
        if size > MAX_FILE_SIZE:
            return f"File too large ({size:,} bytes > {MAX_FILE_SIZE:,} byte limit): {path}"

        suffix = target.suffix.lower()

        if suffix in TEXT_EXTENSIONS:
            try:
                content = target.read_text(encoding="utf-8", errors="replace")
                if len(content) > MAX_TEXT_CHARS:
                    content = content[:MAX_TEXT_CHARS] + f"\n… (truncated at {MAX_TEXT_CHARS:,} chars)"
                return f"Contents of {path} ({size:,} bytes):\n\n{content}"
            except Exception as e:
                return f"Error reading {path}: {e}"

        if suffix in PDF_EXTENSIONS:
            content = _read_pdf(target)
            if len(content) > MAX_TEXT_CHARS:
                content = content[:MAX_TEXT_CHARS] + f"\n… (truncated at {MAX_TEXT_CHARS:,} chars)"
            return f"PDF contents of {path} ({size:,} bytes):\n\n{content}"

        if suffix in WORD_EXTENSIONS:
            content = _read_docx(target)
            if len(content) > MAX_TEXT_CHARS:
                content = content[:MAX_TEXT_CHARS] + f"\n… (truncated at {MAX_TEXT_CHARS:,} chars)"
            return f"Word document contents of {path} ({size:,} bytes):\n\n{content}"

        if suffix in EXCEL_EXTENSIONS:
            content = _read_excel(target)
            if len(content) > MAX_TEXT_CHARS:
                content = content[:MAX_TEXT_CHARS] + f"\n… (truncated at {MAX_TEXT_CHARS:,} chars)"
            return f"Excel contents of {path} ({size:,} bytes):\n\n{content}"

        return f"Unsupported file type ({suffix}): {path}. Supported: text/code, .pdf, .docx, .xlsx"

    return await asyncio.to_thread(_read)


FILE_SKILLS = {
    "list_local_files": list_local_files,
    "read_local_file": read_local_file,
}
